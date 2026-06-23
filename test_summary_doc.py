"""
Test for core/summary_doc.py — the "Download summary" docx generator.
Covers: normal gauge results, not_yet_applicable state, insufficient-
history state, overall document validity, and image embedding (gauge
bars + yield outlook detail chart) via kaleido.
"""
import datetime as dt
import io

from docx import Document
from core.summary_doc import build_summary_docx
from core.dashboard_metrics import MetricResult
from gauge_utils import make_gauge_bar_figure, make_yield_projection_figure

print("=== Test 1: mixed states (normal, not_yet_applicable), no images ===")
water_res = MetricResult("Fallow water gain", "mm", 76.8, 73.1, 26, {}, 2026, None)
nitrogen_res = MetricResult("Fallow nitrogen gain", "kg/ha", 28.3, 76.9, 26, {}, 2026, None)
rain_res = MetricResult("In-crop rain", "mm", None, None, 0, {}, 2026, None, not_yet_applicable=True)
ptq_res = MetricResult("Photothermal quotient", "Cum MJ/m²/°C", None, None, 0, {}, 2026, None, not_yet_applicable=True)
yield_res = MetricResult("Crop yield outlook", "kg/ha", None, None, 0, {}, 2026, None, not_yet_applicable=True)

doc_bytes, warnings = build_summary_docx(
    today=dt.date(2026, 6, 17),
    station_name="DALBY POST OFFICE [QLD] (-27.184, 151.264)",
    soil_name="Deep clay loam",
    fallow_start_md=(1, 1), plant_md=(3, 1), harvest_md=(11, 1),
    wue=25, threshold_water=120,
    gauge_results=[water_res, nitrogen_res, rain_res, ptq_res, yield_res],
)
assert isinstance(doc_bytes, bytes) and len(doc_bytes) > 1000
assert warnings == []
doc = Document(io.BytesIO(doc_bytes))
all_text = "\n".join(p.text for p in doc.paragraphs)
assert "Season tracker summary" in all_text
assert "17 June 2026" in all_text
assert len(doc.inline_shapes) == 0, "No images requested, none should be embedded"
print(f"Generated {len(doc_bytes)} bytes, no warnings, no images (as expected since none were passed).")

results_table = doc.tables[1]
rows_text = [[c.text for c in row.cells] for row in results_table.rows]
assert rows_text[1] == ["Fallow water gain", "77 mm", "above average (26 comparable years)"]
assert rows_text[3] == ["In-crop rain", "—", "Not yet relevant (before plant date)"]
print("Normal and not_yet_applicable rows render correctly.\n")

print("=== Test 2: insufficient-history state ===")
sparse_res = MetricResult("In-crop rain", "mm", 45.0, None, 1, {}, 2026, None)
doc_bytes2, warnings2 = build_summary_docx(
    today=dt.date(2026, 4, 1),
    station_name="Test Station",
    soil_name="Average clay loam",
    fallow_start_md=(1, 1), plant_md=(3, 1), harvest_md=(11, 1),
    wue=25, threshold_water=120,
    gauge_results=[sparse_res],
)
doc2 = Document(io.BytesIO(doc_bytes2))
rows_text2 = [[c.text for c in row.cells] for row in doc2.tables[1].rows]
assert "need more history" in rows_text2[1][2]
print(f"Insufficient-history row: {rows_text2[1]}")
print("Insufficient-history state renders correctly.\n")

print("=== Test 3: WITH images (gauge bars + yield detail chart) ===")
all_normal = [
    MetricResult("Fallow water gain", "mm", 50.0, 60.0, 20, {}, 2026, None),
    MetricResult("Fallow nitrogen gain", "kg/ha", 15.0, 40.0, 20, {}, 2026, None),
    MetricResult("In-crop rain", "mm", 200.0, 80.0, 20, {}, 2026, None),
    MetricResult("Photothermal quotient", "Cum MJ/m²/°C", 60.0, 50.0, 20, {}, 2026, None),
    MetricResult("Crop yield outlook", "kg/ha", 3000.0, 65.0, 20, {}, 2026, None),
]
colors = [("#c9c9c9", "#1f6fb4"), ("#cfe8c9", "#256b1f"), ("#dbeeff", "#0b3d6b"),
          ("#e8d9b0", "#e0a200"), ("#eafce6", "#1e7a1e")]
gauge_figs = [make_gauge_bar_figure(r, lo, hi) for r, (lo, hi) in zip(all_normal, colors)]

# Build a minimal YieldProjection-like object for the chart (reuse the
# dataclass directly so make_yield_projection_figure has real data)
from core.dashboard_metrics import YieldProjection
import pandas as pd
import numpy as np

dates_range = pd.date_range("2026-03-01", "2026-11-01", freq="D")
n = len(dates_range)
p10 = pd.Series(np.linspace(0, 3000, n), index=dates_range)
p50 = pd.Series(np.linspace(0, 5000, n), index=dates_range)
p90 = pd.Series(np.linspace(0, 8000, n), index=dates_range)
actual = pd.Series(np.linspace(0, 1200, 90), index=dates_range[:90])
projected = pd.Series(np.linspace(1200, 4500, n - 89), index=dates_range[89:])
yield_proj = YieldProjection(
    dates=dates_range, p10=p10, p50=p50, p90=p90, actual=actual, projected=projected,
    cond_p10=pd.Series(np.linspace(1200, 3000, n - 89), index=dates_range[89:]),
    cond_p50=pd.Series(np.linspace(1200, 4500, n - 89), index=dates_range[89:]),
    cond_p90=pd.Series(np.linspace(1200, 6000, n - 89), index=dates_range[89:]),
    today=dt.date(2026, 5, 30), harvest_date=dt.date(2026, 11, 1), n_comparable_years=20,
)
yield_fig = make_yield_projection_figure(yield_proj)

doc_bytes3, warnings3 = build_summary_docx(
    today=dt.date(2026, 8, 1),
    station_name="Emerald",
    soil_name="Average sand loam",
    fallow_start_md=(1, 1), plant_md=(3, 1), harvest_md=(11, 1),
    wue=22, threshold_water=110,
    gauge_results=all_normal,
    gauge_figures=gauge_figs,
    yield_figure=yield_fig,
)
print(f"Generated {len(doc_bytes3)} bytes, warnings = {warnings3}")
assert len(doc_bytes3) > 50000, "Document with 6 embedded images should be substantially larger"

doc3 = Document(io.BytesIO(doc_bytes3))
n_images = len(doc3.inline_shapes)
print(f"Embedded images: {n_images}")
assert n_images == 6, f"Expected 6 images (5 gauges + 1 yield chart), got {n_images}"

rows_text3 = [[c.text for c in row.cells] for row in doc3.tables[1].rows]
expected_descriptors = ["above average", "below average", "high", "average", "above average"]
for i, expected in enumerate(expected_descriptors):
    assert expected in rows_text3[i + 1][2], f"Row {i+1}: expected '{expected}' in {rows_text3[i+1][2]}"

all_text3 = "\n".join(p.text for p in doc3.paragraphs)
assert "Crop yield outlook — detail" in all_text3
print("All 6 images embedded, results table correct, yield detail heading present.\n")

print("All summary_doc tests passed.")
