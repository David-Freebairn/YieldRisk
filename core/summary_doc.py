"""
core/summary_doc.py

Builds the "Download summary" Word document: setup inputs (site, soil,
dates, WUE, threshold water), the five gauge results (value, unit, band
descriptor) for the current season, the five gauge bars as images, and
the detailed yield outlook chart (10-90%ile plume) as an image.

Uses python-docx (a Python library, distinct from the docx-js/Node
tooling used for Claude's own document-authoring skill) since this
document is generated at runtime inside the running Streamlit app from
live MetricResult data, not authored ahead of time as a static file.

Plotly figures are exported to static PNG bytes via kaleido (fig.to_image)
before embedding, since Word documents can't contain interactive charts.
"""

import io
import time
import datetime as dt

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from gauge_utils import percentile_to_descriptor

BRAND_BLUE = RGBColor(0x1F, 0x6F, 0xB4)
DARK_GREY = RGBColor(0x33, 0x33, 0x33)

GAUGE_IMAGE_WIDTH_PX = 900
GAUGE_IMAGE_HEIGHT_PX = 110
YIELD_IMAGE_WIDTH_PX = 1000
YIELD_IMAGE_HEIGHT_PX = 480


def _add_heading(doc, text, size=18, color=BRAND_BLUE, bold=True, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_kv_row(table, label, value):
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value
    row[0].paragraphs[0].runs[0].font.bold = True
    row[0].paragraphs[0].runs[0].font.size = Pt(10)
    row[1].paragraphs[0].runs[0].font.size = Pt(10)


def _fig_to_png_bytes(fig, width_px, height_px, retries=2):
    """Export a Plotly figure to PNG bytes via kaleido.

    Returns (png_bytes, error_message). On success, error_message is None.
    On failure, png_bytes is None and error_message describes what went
    wrong — most commonly kaleido needing a Chrome/Chromium install it
    can't find (kaleido>=1.0 expects a system browser; kaleido==0.2.x
    bundles its own, which is what this project pins in requirements.txt
    to avoid that extra dependency). The caller decides how to surface
    this rather than the chart simply vanishing with no explanation.

    Retries a couple of times with a short pause before giving up — kaleido
    launches a headless browser subprocess per export, and that can fail
    transiently (e.g. on the very first export in a process, or under
    momentary resource pressure) even when it would succeed a moment later.
    Larger, more complex figures (more traces, bigger canvas) are more
    likely to hit this than small simple ones, which is consistent with
    five small gauge bars succeeding while a 13-trace detail chart fails.
    """
    last_error = None
    for attempt in range(retries + 1):
        try:
            return fig.to_image(format="png", width=width_px, height=height_px, scale=2), None
        except Exception as e:
            last_error = f"{type(e).__name__}: {e}"
            if attempt < retries:
                time.sleep(0.75)
    return None, last_error


def build_summary_docx(
    today: dt.date,
    station_name: str,
    soil_name: str,
    fallow_start_md, plant_md, harvest_md,
    wue: float, threshold_water: float,
    gauge_results: list,
    gauge_figures: list = None,
    yield_figure=None,
) -> tuple:
    """
    Build the season summary Word document. Returns (docx_bytes, warnings)
    where warnings is a list of strings describing any charts that
    couldn't be rendered as images (e.g. if kaleido/Chrome isn't available)
    — the document is still generated and returned even if some or all
    images fail, just without those particular pictures, so the caller can
    decide whether/how to tell the user.

    gauge_results: list of core.dashboard_metrics.MetricResult objects,
    in display order (water, nitrogen, rain, ptq, yield).
    gauge_figures: optional list of Plotly figures (from
    gauge_utils.make_gauge_bar_figure), same order as gauge_results — the
    five gauge bars as shown on screen, embedded as images below the
    results table. If None, gauge images are skipped (table-only, as in
    earlier versions of this document).
    yield_figure: optional Plotly figure (from
    gauge_utils.make_yield_projection_figure) — the detailed yield outlook
    10-90%ile plume chart, embedded as an image at the end of the
    document. If None, it's skipped.
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    content_width_in = 6.5  # 8.5" page - 1" left - 1" right margin

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_heading(doc, "Yieldrisk summary", size=22, space_after=2)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"As of {today.strftime('%d %B %Y')}  ·  {station_name}")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = DARK_GREY
    sub.paragraph_format.space_after = Pt(14)

    month_names = ["Jan", "Feb", "Mar", "Apr", "May", "Jun",
                   "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

    def fmt_md(md):
        d, m = md[1], md[0]
        return f"{d} {month_names[m - 1]}"

    _add_heading(doc, "Setup", size=14, space_after=4)
    setup_table = doc.add_table(rows=0, cols=2)
    setup_table.autofit = True
    _add_kv_row(setup_table, "Site", station_name)
    _add_kv_row(setup_table, "Soil type", soil_name)
    _add_kv_row(setup_table, "Fallow start", fmt_md(fallow_start_md))
    _add_kv_row(setup_table, "Plant date", fmt_md(plant_md))
    _add_kv_row(setup_table, "Harvest date", fmt_md(harvest_md))
    _add_kv_row(setup_table, "WUE", f"{wue:.0f} kg/ha/mm")
    _add_kv_row(setup_table, "Threshold water", f"{threshold_water:.0f} mm")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    _add_heading(doc, "How are we going", size=14, space_after=4)
    results_table = doc.add_table(rows=0, cols=3)
    results_table.autofit = True
    header_row = results_table.add_row().cells
    for i, h in enumerate(["Metric", "Value", "vs. history"]):
        header_row[i].text = h
        header_row[i].paragraphs[0].runs[0].font.bold = True
        header_row[i].paragraphs[0].runs[0].font.size = Pt(10)

    for result in gauge_results:
        row = results_table.add_row().cells
        row[0].text = result.label
        if result.not_yet_applicable:
            row[1].text = "—"
            row[2].text = "Not yet relevant (before plant date)"
        elif result.current_value is None:
            row[1].text = "No data"
            row[2].text = "—"
        else:
            row[1].text = f"{result.current_value:.0f} {result.unit}"
            descriptor = percentile_to_descriptor(result.percentile)
            if descriptor is not None:
                row[2].text = f"{descriptor} ({result.n_comparable_years} comparable years)"
            else:
                row[2].text = f"Only {result.n_comparable_years} comparable years — need more history"
        for cell in row:
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    # ── Gauge bar images (as shown on screen) ───────────────────────────
    image_warnings = []
    if gauge_figures:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        for fig, result in zip(gauge_figures, gauge_results):
            if fig is None:
                continue
            png_bytes, err = _fig_to_png_bytes(fig, GAUGE_IMAGE_WIDTH_PX, GAUGE_IMAGE_HEIGHT_PX)
            if png_bytes is None:
                image_warnings.append(f"{result.label} gauge: {err}")
                continue
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(content_width_in))
            pic_para = doc.paragraphs[-1]
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_para.paragraph_format.space_after = Pt(2)

    # ── Detailed yield outlook chart ─────────────────────────────────────
    if yield_figure is not None:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        _add_heading(doc, "Crop yield outlook — detail", size=14, space_after=4)
        png_bytes, err = _fig_to_png_bytes(yield_figure, YIELD_IMAGE_WIDTH_PX, YIELD_IMAGE_HEIGHT_PX)
        if png_bytes is not None:
            doc.add_picture(io.BytesIO(png_bytes), width=Inches(content_width_in))
            pic_para = doc.paragraphs[-1]
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            image_warnings.append(f"Yield outlook chart: {err}")
            note = doc.add_paragraph()
            note_run = note.add_run(
                f"(Yield outlook chart could not be rendered for this document. "
                f"Reason: {err})"
            )
            note_run.italic = True
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Generated by Yieldrisk. Percentile bands compare this "
        "season's simulated/observed values against comparable historical years "
        "(1996 onwards) at this site, soil and dates."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer_run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue(), image_warnings
